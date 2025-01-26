"""
Hybrid Agent Enhanced 

"""

import os
import re
import logging
import requests
from dataclasses import dataclass
from functools import lru_cache
from typing import List, Dict, Optional, Tuple, Any
from datetime import datetime
from pathlib import Path

# -----------------------------
# Optional LlamaIndex (GPT Index) imports
# -----------------------------
try:
    from llama_index import (
        VectorStoreIndex,
        SQLStructStoreIndex,
        RouterRetriever,
        ServiceContext,
        RouterPrompt
    )
    from llama_index.llms import OpenAI as LlamaIndexOpenAI
    LLAMA_INDEX_AVAILABLE = True
except ImportError:
    LLAMA_INDEX_AVAILABLE = False

# -----------------------------
# LangChain / Community imports
# -----------------------------
from langchain_community.utilities import SQLDatabase
from langchain_community.agent_toolkits.sql.toolkit import SQLDatabaseToolkit
from langchain_community.agent_toolkits import create_sql_agent
from langchain_community.vectorstores import FAISS
from langchain_community.retrievers import BM25Retriever

from langchain.retrievers import EnsembleRetriever
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

# Because some LangChain versions differ in pydantic usage
from langchain_core.pydantic_v1 import BaseModel

# Instead of standard LC chat_models/embeddings, we use these:
from langchain_community.chat_models import ChatOpenAI
from langchain_community.embeddings import OpenAIEmbeddings

# Document processing
import PyPDF2
import docx

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[logging.FileHandler("hybrid_agent.log"), logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

########################################################################
# 1) CONFIGURATION
########################################################################

class Config(BaseModel):
    """
    Configuration for the Hybrid Agent system.
    
    Adjust the LLM_MODEL to "gpt-3.5-turbo" or "gpt-4" (or other).
    Increase TEMPERATURE if you want more creative/human-like outputs.
    """
    OPENAI_API_KEY: str = os.getenv("OPENAI_API_KEY", "")
    DEEPSEEK_API_KEY: str = os.getenv("DEEPSEEK_API_KEY", "")
    # e.g. "gpt-4", "gpt-3.5-turbo", or any advanced "O1" model name
    LLM_MODEL: str = "gpt-3.5-turbo"
    TEMPERATURE: float = 0.7
    CHUNK_SIZE: int = 500
    CHUNK_OVERLAP: int = 100
    CACHE_SIZE: int = 100
    SQL_CONFIDENCE_THRESHOLD: float = 0.6
    RAG_CONFIDENCE_THRESHOLD: float = 0.4
    DENSE_WEIGHT: float = 0.7
    SPARSE_WEIGHT: float = 0.3
    ENABLE_CACHE: bool = True
    CACHE_TTL: int = 3600

########################################################################
# 2) DATA MODELS
########################################################################

@dataclass
class QueryResult:
    """
    Represents the result of a query, including:
    - content: the textual answer
    - confidence: numeric confidence in the answer
    - source: 'sql', 'rag', 'greeting', 'guardrail', 'error', etc.
    - metadata: any extra info
    - timestamp: when the answer was produced
    """
    content: str
    confidence: float
    source: str
    metadata: Dict
    timestamp: datetime = datetime.now()

@dataclass
class SchemaInfo:
    """Database schema information."""
    tables: Dict[str, List[str]]
    relationships: Dict[str, List[str]]

@dataclass
class ReasoningResult:
    """
    Result from the reasoning model (DeepSeek or fallback).
    - understood_query: paraphrased or clarified query
    - confidence: how confident the system is in that analysis
    - sub_queries: if the user asked multiple things, we can split them
    - context: extra analysis context
    """
    understood_query: str
    confidence: float
    sub_queries: List[str]
    context: Dict[str, Any]

@dataclass
class RoutingDecision:
    """
    Decides the route: 'sql', 'rag', 'greeting', 'guardrail', etc.
    - confidence: numeric confidence
    - strategy: 'pattern', 'semantic', 'hybrid', or 'fallback'
    - context: additional details if needed
    """
    route: str
    confidence: float
    strategy: str
    context: Dict[str, Any]

########################################################################
# 3) DOCUMENT PROCESSING
########################################################################

def load_documents(docs_path: Path, config: Config) -> List[Document]:
    """
    Loads all PDF and DOCX files from the given directory,
    then splits them into smaller chunks for RAG.
    """
    documents = []
    for file_path in docs_path.glob("*.*"):
        try:
            if file_path.suffix.lower() == '.pdf':
                doc = load_pdf(file_path)
            elif file_path.suffix.lower() == '.docx':
                doc = load_docx(file_path)
            else:
                continue

            if doc:
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=config.CHUNK_SIZE,
                    chunk_overlap=config.CHUNK_OVERLAP
                )
                split_docs = text_splitter.split_documents([doc])
                documents.extend(split_docs)

        except Exception as e:
            logger.error(f"Error loading document {file_path}: {e}")

    if not documents:
        logger.warning(f"No documents loaded from {docs_path} (PDF/DOCX).")
    return documents

def load_pdf(file_path: Path) -> Optional[Document]:
    """Loads a single PDF file as a Document."""
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            content = ' '.join(page.extract_text() or "" for page in reader.pages)
            return Document(
                page_content=content,
                metadata={'source': file_path.name, 'type': 'pdf'}
            )
    except Exception as e:
        logger.error(f"Error loading PDF {file_path}: {e}")
        return None

def load_docx(file_path: Path) -> Optional[Document]:
    """Loads a single DOCX file as a Document."""
    try:
        doc = docx.Document(file_path)
        content = ' '.join(paragraph.text for paragraph in doc.paragraphs)
        return Document(
            page_content=content,
            metadata={'source': file_path.name, 'type': 'docx'}
        )
    except Exception as e:
        logger.error(f"Error loading DOCX {file_path}: {e}")
        return None

########################################################################
# 4) REASONING MODEL (DEEPSEEK OR FALLBACK)
########################################################################

class ReasoningModel:
    """
    This model attempts to interpret the user's query semantically using DeepSeek
    or a fallback approach if that fails. It’s primarily used for semantic
    confidence scoring in the SmartRouter. The sub-query breakdown is now
    handled by the QueryDecomposer.
    """
    def __init__(self, config: Config):
        self.config = config
        self.deepseek_api_key = config.DEEPSEEK_API_KEY
        self.logger = logging.getLogger(__name__)

    def understand_query(self, query: str) -> ReasoningResult:
        """
        If DEEPSEEK_API_KEY is available, do an external call to interpret the query.
        Otherwise, do a simple fallback.
        """
        if not self.deepseek_api_key:
            # If no DeepSeek key, just fallback immediately
            return self._fallback_analysis(query)

        try:
            headers = {"Authorization": f"Bearer {self.deepseek_api_key}"}
            response = requests.post(
                "https://api.deepseek.com/v1/comprehend",
                json={"query": query},
                headers=headers
            )

            if response.status_code == 200:
                analysis = response.json()
                return ReasoningResult(
                    understood_query=analysis.get('understood_query', query),
                    confidence=analysis.get('confidence', 0.8),
                    sub_queries=[query],  # We now handle sub-query decomposition separately
                    context=analysis.get('context', {})
                )
            else:
                return self._fallback_analysis(query)
        except Exception as e:
            self.logger.error(f"DeepSeek error: {e}")
            return self._fallback_analysis(query)

    def _fallback_analysis(self, query: str) -> ReasoningResult:
        """
        Minimal fallback if DeepSeek is unavailable.
        """
        return ReasoningResult(
            understood_query=query,
            confidence=0.6,
            sub_queries=[query],
            context={}
        )

########################################################################
# 5) GUARDRAILS
########################################################################

class GuardrailChecker:
    """
    Simple guardrail system that looks for disallowed or sensitive queries.
    If matched, the system should refuse to answer.
    """
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        # Example patterns: personal data, illegal instructions, etc.
        self.disallowed_patterns = [
            r"(\bssn\b|\bsocial security\b)",
            r"(\bcredit card\b)",
            r"(bomb making|terrorist|child abuse)",
        ]
    
    def check(self, query: str) -> Optional[str]:
        """Returns a string describing the violation if matched, else None."""
        low_query = query.lower()
        for pattern in self.disallowed_patterns:
            if re.search(pattern, low_query):
                return f"Matched guardrail pattern: {pattern}"
        return None

########################################################################
# 6) SPELLING CORRECTION (LLM-BASED)
########################################################################

class SpellCorrector:
    """
    Uses an LLM at temperature=0.0 to correct user spelling.
    You could swap in a local library (like 'pyspellchecker') if desired.
    """
    def __init__(self, config: Config):
        self.config = config
        self.llm = ChatOpenAI(
            model=config.LLM_MODEL,
            temperature=0.0,  # minimal creative changes
            api_key=config.OPENAI_API_KEY
        )
        self.logger = logging.getLogger(__name__)

    def correct(self, query: str) -> str:
        """Attempt to correct spelling errors in the query."""
        try:
            prompt = f"""
You are a helpful assistant. The user query below may contain spelling mistakes.
Please correct any spelling mistakes, but keep the original meaning.

Query: "{query}"
Corrected query:
"""
            response = self.llm.invoke(prompt)
            return response.content.strip()
        except Exception as e:
            self.logger.error(f"Spell correction error: {e}")
            return query  # fallback: return original if error

########################################################################
# 7) ANSWER REFINER
########################################################################

class AnswerRefiner:
    """
    Takes a raw answer (from SQL or RAG) and re-writes it in a more
    human-like or conversational style, using a higher temperature if desired.
    """
    def __init__(self, config: Config):
        self.config = config
        self.llm = ChatOpenAI(
            model=config.LLM_MODEL,
            temperature=config.TEMPERATURE,
            api_key=config.OPENAI_API_KEY
        )
        self.logger = logging.getLogger(__name__)

    def refine(self, raw_answer: str) -> str:
        """Rewrite the raw answer in a friendlier, more natural style."""
        try:
            prompt = f"""
You are a friendly, conversational assistant. 
Rewrite the following answer so it sounds natural and human-like, but do not alter the factual content:

Original answer:
{raw_answer}

Rewritten answer:
"""
            response = self.llm.invoke(prompt)
            return response.content.strip()
        except Exception as e:
            self.logger.error(f"Answer refinement error: {e}")
            return raw_answer

########################################################################
# 8) RETRIEVERS: SQL + RAG
########################################################################

class SQLRetriever:
    """
    SQL database retriever via a LangChain-based LLM agent.
    Attempts to transform natural language into SQL queries.
    """
    def __init__(self, db: SQLDatabase, config: Config, refiner: AnswerRefiner):
        self.db = db
        self.config = config
        self.logger = logging.getLogger(__name__)
        self.refiner = refiner

        self.llm = ChatOpenAI(
            model=config.LLM_MODEL,
            temperature=config.TEMPERATURE,
            api_key=config.OPENAI_API_KEY
        )
        toolkit = SQLDatabaseToolkit(db=db, llm=self.llm)
        self.agent = create_sql_agent(
            llm=self.llm,
            toolkit=toolkit,
            agent_type="openai-tools",
            verbose=True
        )

    def retrieve(self, query: str) -> QueryResult:
        """
        Invoke the SQL agent on the user query. If successful,
        refine the final text. Otherwise, return empty content.
        """
        try:
            result = self.agent.invoke(query)
            if isinstance(result, dict):
                content = str(result.get('output', ''))
            else:
                content = str(result)

            # Refine output to be more human-like
            refined = self.refiner.refine(content)

            # If there's actual content, set confidence=1, else 0
            if refined.strip():
                return QueryResult(
                    content=refined,
                    confidence=1.0,
                    source='sql',
                    metadata={'query_type': 'sql'}
                )
            else:
                return QueryResult(
                    content="",
                    confidence=0.0,
                    source='sql',
                    metadata={'error': 'Empty SQL response'}
                )
        except Exception as e:
            self.logger.error(f"SQL retrieval error: {e}")
            return QueryResult(
                content="",
                confidence=0.0,
                source='sql',
                metadata={'error': str(e)}
            )

class RAGRetriever:
    """
    Retrieval-Augmented Generation using an ensemble retriever:
    - Dense (FAISS) weighted
    - Sparse (BM25) weighted
    Summarizes the top docs to produce an answer.
    """
    def __init__(self, documents: List[Document], config: Config, refiner: AnswerRefiner):
        self.config = config
        self.documents = documents
        self.refiner = refiner
        self.logger = logging.getLogger(__name__)

        # Initialize sub-retrievers
        self.dense_retriever = self._init_dense_retriever()
        self.sparse_retriever = self._init_sparse_retriever()
        self.ensemble_retriever = self._init_ensemble_retriever()

        self.llm = ChatOpenAI(
            model=config.LLM_MODEL,
            temperature=config.TEMPERATURE,
            api_key=config.OPENAI_API_KEY
        )

    @lru_cache(maxsize=None)
    def _init_dense_retriever(self):
        embeddings = OpenAIEmbeddings(api_key=self.config.OPENAI_API_KEY)
        return FAISS.from_documents(self.documents, embeddings).as_retriever(
            search_kwargs={"k": 5}
        )

    def _init_sparse_retriever(self):
        return BM25Retriever.from_documents(self.documents, k=5)

    def _init_ensemble_retriever(self):
        return EnsembleRetriever(
            retrievers=[self.sparse_retriever, self.dense_retriever],
            weights=[self.config.SPARSE_WEIGHT, self.config.DENSE_WEIGHT]
        )

    def retrieve(self, query: str) -> QueryResult:
        """
        Use an ensemble of sparse/dense retrieval, then summarize
        the top documents using the LLM, and refine the answer.
        """
        try:
            docs = self.ensemble_retriever.invoke(query)
            if not docs:
                return QueryResult(
                    content="",
                    confidence=0.0,
                    source='rag',
                    metadata={'found_docs': 0}
                )

            # Summarize the top docs
            summary = self._summarize_docs(query, docs)
            refined = self.refiner.refine(summary)

            # If there's content, set confidence=0.8
            if refined.strip():
                return QueryResult(
                    content=refined,
                    confidence=0.8,
                    source='rag',
                    metadata={'found_docs': len(docs)}
                )
            else:
                return QueryResult(
                    content="",
                    confidence=0.0,
                    source='rag',
                    metadata={'found_docs': len(docs), 'error': 'Empty RAG summary'}
                )
        except Exception as e:
            self.logger.error(f"RAG retrieval error: {e}")
            return QueryResult(
                content="",
                confidence=0.0,
                source='rag',
                metadata={'error': str(e)}
            )

    def _summarize_docs(self, query: str, docs: List[Document]) -> str:
        """
        Pass doc content plus the user query to the LLM for summarization.
        """
        # Combine content from top docs (limit ~3 for prompt size)
        combined = "\n".join(doc.page_content[:3000] for doc in docs[:3])
        prompt = f"""
You are a helpful assistant. Summarize the following text to answer the query: "{query}"

Text:
{combined}

Please provide a concise answer:
"""
        response = self.llm.invoke(prompt)
        return response.content

########################################################################
# 9) ROUTERS: QUERY ROUTER + SMART ROUTER
########################################################################

class QueryRouter:
    """
    Basic regex-based router that checks for greetings, known SQL patterns,
    or RAG patterns. If no patterns match, defaults to RAG.
    """
    def __init__(self, schema_info: SchemaInfo, config: Config):
        self.schema_info = schema_info
        self.config = config
        # Extended to include "composer", "track", "artist", etc.
        self.sql_patterns = {
            'aggregation': r'\b(count|sum|average|avg|total|mean|most|maximum)\b',
            'comparison': r'\b(more than|less than|greater|highest|lowest|most)\b',
            'temporal': r'\b(when|date|year|month|time|period)\b',
            'numerical': r'\b(how many|how much|price|cost|revenue|sales)\b',
            'music': r'\b(artist|track|album|song|composer|composers)\b',
            'customer': r'\b(customers?|client|buyer|purchase|spending)\b',
            'genre': r'\b(genre|category|type|style)\b'
        }
        self.rag_patterns = {
            'conceptual': r'\b(what is|explain|describe|how does|why)\b',
            'research': r'\b(paper|research|study|analysis|findings)\b',
            'procedural': r'\b(how to|steps|process|procedure)\b'
        }
        self.greeting_patterns = {
            'hello': r'\b(hello|hi|hey|greetings)\b',
            'goodbye': r'\b(goodbye|bye|see you|farewell)\b',
            'thanks': r'\b(thank you|thanks|appreciate)\b',
            'wellbeing': r'\b(how are you|hows it going|whatsup)\b'
        }

    def route_query(self, query: str) -> Tuple[str, float]:
        """
        Return (route, confidence).
        route can be: 'greeting', 'sql', or 'rag'.
        """
        query_lower = query.lower()

        # 1) Check greetings
        for pattern in self.greeting_patterns.values():
            if re.search(pattern, query_lower):
                return 'greeting', 1.0

        # 2) Check if user references DB schema strongly
        schema_score = self._check_schema_match(query_lower)
        if schema_score > self.config.SQL_CONFIDENCE_THRESHOLD:
            return 'sql', schema_score

        # 3) Pattern-based scoring
        sql_score = self._check_patterns(query_lower, self.sql_patterns)
        rag_score = self._check_patterns(query_lower, self.rag_patterns)

        # If no patterns match, default to RAG
        if sql_score == 0 and rag_score == 0:
            return 'rag', 0.0

        total_score = sql_score + rag_score
        if total_score == 0:
            return 'rag', 0.0

        # Whichever is higher gets the route
        if sql_score > rag_score:
            return 'sql', sql_score / total_score
        else:
            return 'rag', rag_score / total_score

    def _check_schema_match(self, query: str) -> float:
        """
        If the user references actual table or column names,
        we raise the score for SQL. 
        """
        score = 0.0
        for table_name, columns in self.schema_info.tables.items():
            if table_name.lower() in query:
                score += 0.5
            for col in columns:
                if col.lower() in query:
                    score += 0.3
        return min(score, 1.0)

    def _check_patterns(self, query: str, patterns: dict) -> float:
        """
        For each pattern set, if found in the query, add to the score.
        e.g. 'composer' or 'most' for SQL, 'what is' for RAG, etc.
        """
        score = 0.0
        for pattern in patterns.values():
            if re.search(pattern, query):
                score += 1.0
        return score

class SmartRouter:
    """
    Combines pattern-based logic with optional semantic confidence (ReasoningModel).
    If the user references DB schema strongly, we route to SQL directly.
    Otherwise, we combine the pattern confidence with the semantic confidence.
    """
    def __init__(self, schema_info: SchemaInfo, config: Config, reasoning_model: ReasoningModel):
        self.schema_info = schema_info
        self.config = config
        self.reasoning_model = reasoning_model
        self.pattern_router = QueryRouter(schema_info, config)
        self.logger = logging.getLogger(__name__)

    def route_query(self, query: str) -> RoutingDecision:
        """
        Return a RoutingDecision with route='sql','rag','greeting', 
        plus confidence and strategy.
        """
        try:
            # 1) Semantic analysis (DeepSeek or fallback)
            understanding = self.reasoning_model.understand_query(query)

            # 2) Pattern-based routing
            pattern_route, pattern_confidence = self.pattern_router.route_query(query)

            if pattern_route == 'greeting':
                return RoutingDecision('greeting', 1.0, 'pattern', {})

            # 3) Check for strong DB schema references
            schema_confidence = self._check_schema_relevance(understanding.understood_query)
            if schema_confidence > self.config.SQL_CONFIDENCE_THRESHOLD:
                return RoutingDecision('sql', schema_confidence, 'schema', {
                    'understanding': understanding.context
                })

            # 4) Weighted combo of semantic & pattern
            # (e.g. 70% from DeepSeek, 30% from pattern)
            semantic_weight = 0.7
            pattern_weight = 0.3
            final_confidence = (
                semantic_weight * understanding.confidence +
                pattern_weight * pattern_confidence
            )

            return RoutingDecision(
                route=pattern_route,
                confidence=final_confidence,
                strategy='hybrid',
                context={'understanding': understanding.context}
            )
        except Exception as e:
            self.logger.error(f"Error in SmartRouter: {e}")
            # fallback to RAG
            return RoutingDecision('rag', 0.5, 'fallback', {'error': str(e)})

    def _check_schema_relevance(self, query: str) -> float:
        """Same logic as pattern router, but for the understood query text."""
        score = 0.0
        for table_name, columns in self.schema_info.tables.items():
            if table_name.lower() in query.lower():
                score += 0.5
            for col in columns:
                if col.lower() in query.lower():
                    score += 0.3
        return min(score, 1.0)

########################################################################
# 10) QUERY DECOMPOSER
########################################################################

class QueryDecomposer:
    """
    Break down complex user queries into sub-queries using an LLM prompt.
    """
    def __init__(self, config: Config):
        self.config = config
        self.llm = ChatOpenAI(
            model=config.LLM_MODEL,
            temperature=config.TEMPERATURE,
            api_key=config.OPENAI_API_KEY
        )

    def decompose(self, query: str) -> List[str]:
        """
        A prompt that tries to produce multiple sub-queries only if
        there are truly multiple distinct requests.
        """
        prompt = f"""
You are a helpful assistant. The user query may contain multiple or combined questions or statements.

- If the query only has a greeting plus a single question, treat it as a single sub-query.
- If the query has multiple, unrelated questions, break them into sub-queries.
- If there is only one question/statement, produce exactly one sub-query.

User query: "{query}"

Sub-queries:
"""
        response = self.llm.invoke(prompt)
        # Split by line breaks, ignoring empty lines
        raw_lines = [line.strip() for line in response.content.split("\n")]
        sub_queries = [line for line in raw_lines if line]

        return sub_queries


########################################################################
# 11) MAIN HYBRID AGENT
########################################################################

class HybridAgent:
    """
    The main agent that ties everything together:
    - Connects to DB
    - Loads docs
    - Sets up RAG & SQL retrievers
    - Uses a SmartRouter to decide the route
    - Has guardrails & spell correction
    - Falls back from SQL -> RAG -> "no answer" if needed
    - Optionally includes a LlamaIndex RouterRetriever
    - NEW: Actively uses QueryDecomposer to break multi-part queries.
    """
    def __init__(self, database_uri: str, docs_directory: str, config: Optional[Config] = None):
        self.config = config or Config()
        self.logger = logging.getLogger(__name__)

        # 1) Connect to DB
        try:
            self.db = SQLDatabase.from_uri(database_uri)
            self.logger.info("Database connection successful.")
            self._log_available_tables()
        except Exception as e:
            self.logger.error(f"Database connection failed: {e}")
            raise ConnectionError(f"Failed to connect to database: {e}")

        # 2) Load documents for RAG
        docs_path = Path(docs_directory)
        if not docs_path.exists():
            raise ValueError(f"Documents directory does not exist: {docs_directory}")

        self.documents = load_documents(docs_path, self.config)

        # 3) Extract DB schema info
        self.schema_info = self._extract_schema_info()

        # 4) Initialize core components
        self.reasoning_model = ReasoningModel(self.config)
        self.smart_router = SmartRouter(self.schema_info, self.config, self.reasoning_model)
        self.guardrail_checker = GuardrailChecker()
        self.spell_corrector = SpellCorrector(self.config)
        self.answer_refiner = AnswerRefiner(self.config)
        self.sql_retriever = SQLRetriever(self.db, self.config, self.answer_refiner)
        self.rag_retriever = RAGRetriever(self.documents, self.config, self.answer_refiner)
        self.decomposer = QueryDecomposer(self.config)

        # 5) Optionally build a LlamaIndex router
        self.llamaindex_router_retriever = None
        self._build_llamaindex_router()  # no-op if not installed

    def _log_available_tables(self):
        """Log the list of tables in the connected DB."""
        try:
            table_names = self.db.get_usable_table_names()
            self.logger.info(f"Available tables: {table_names}")
        except Exception as e:
            self.logger.error(f"Error fetching table names: {e}")

    def _extract_schema_info(self) -> SchemaInfo:
        """
        Builds a SchemaInfo from the DB. We read each table name + its columns.
        """
        try:
            tables = {}
            relationships = {}

            table_names = self.db.get_usable_table_names()
            for table_name in table_names:
                try:
                    table_info = self.db.get_table_info(table_name)
                    columns = []
                    for field_info in table_info:
                        if hasattr(field_info, 'name'):
                            column_name = field_info.name
                        elif hasattr(field_info, 'key'):
                            column_name = field_info.key
                        else:
                            continue
                        columns.append(column_name)
                    tables[table_name] = columns
                    relationships[table_name] = []
                    self.logger.info(f"Processed table {table_name} with columns: {columns}")
                except Exception as e:
                    self.logger.error(f"Error processing table {table_name}: {e}")
                    continue

            return SchemaInfo(tables=tables, relationships=relationships)
        except Exception as e:
            self.logger.error(f"Error extracting schema info: {e}")
            return SchemaInfo(tables={}, relationships={})

    def _build_llamaindex_router(self):
        """
        Build a RouterRetriever from LlamaIndex if installed.
        This can auto-route between a VectorStoreIndex (docs) and a SQLStructStoreIndex.
        """
        if not LLAMA_INDEX_AVAILABLE:
            self.logger.warning("LlamaIndex not available; cannot build RouterRetriever.")
            return

        try:
            llm = LlamaIndexOpenAI(
                api_key=self.config.OPENAI_API_KEY,
                model=self.config.LLM_MODEL
            )
            service_context = ServiceContext.from_defaults(llm=llm)

            # Convert documents to LlamaIndex "Document" objects
            from llama_index.readers.schema.base import Document as LIDocument
            doc_index = VectorStoreIndex.from_documents(
                [LIDocument(text=d.page_content, extra_info=d.metadata) for d in self.documents],
                service_context=service_context
            )
            sql_index = SQLStructStoreIndex(
                uri=self.db.uri,
                service_context=service_context
            )

            router_prompt = RouterPrompt(
                prompt_template="""
You are a router that decides whether to query the SQL index or the doc index.
If the question requires structured data (numbers, stats, etc.), choose the SQL index.
Otherwise, choose the doc index.

Question: {query_str}
""".strip()
            )

            self.llamaindex_router_retriever = RouterRetriever(
                selector=router_prompt,
                retrievers={
                    "sql_index": sql_index.as_retriever(),
                    "doc_index": doc_index.as_retriever(),
                }
            )
            self.logger.info("LlamaIndex RouterRetriever built successfully.")
        except Exception as e:
            self.logger.error(f"Error building LlamaIndex RouterRetriever: {e}")

    ####################################################################
    # PUBLIC QUERY METHODS
    ####################################################################

    def process_query(self, query: str) -> QueryResult:
        """
        Main method to handle a user query. 
        Steps:
        1) Guardrail check
        2) Spell-correct
        3) Decompose into sub-queries
        4) For each sub-query: route -> retrieve -> fallback -> gather results
        5) Combine results and refine
        """
        try:
            self.logger.info(f"Original user query: {query}")

            # 1) Guardrails
            violation = self.guardrail_checker.check(query)
            if violation:
                return QueryResult(
                    content="I’m sorry, but I can’t help with that request.",
                    confidence=1.0,
                    source='guardrail',
                    metadata={'violation': violation}
                )

            # 2) Spelling correction
            corrected_query = self.spell_corrector.correct(query)
            self.logger.info(f"Spell-corrected query: {corrected_query}")

            # 3) Decompose into sub-queries
            sub_queries = self.decomposer.decompose(corrected_query)

            # -- NEW POST-PROCESSING LOGIC --
            # If we have exactly 2 sub-queries, and the first is just a short greeting,
            # merge them back into a single query.
            if len(sub_queries) == 2 and len(sub_queries[0].split()) <= 3:
                # e.g. sub_queries[0] = "Hey there", sub_queries[1] = "Do you know who are the composers ..."
                merged = sub_queries[0] + " " + sub_queries[1]
                sub_queries = [merged]

            # If decomposition yields nothing, fallback
            if not sub_queries:
                sub_queries = [corrected_query]

            # 4) Process each sub-query
            results_list = []
            for sq in sub_queries:
                sq_result = self._process_sub_query(sq)
                results_list.append((sq, sq_result))

            # Combine results and refine
            combined_text = self._format_results(results_list)
            refined = self.answer_refiner.refine(combined_text)

            # Find min confidence across all sub-queries to give an overall measure
            min_conf = min((res.confidence for (_, res) in results_list), default=1.0)

            # If everything is empty, show fallback
            if all(not res.content for (_, res) in results_list):
                return QueryResult(
                    content="I’m sorry, but I couldn’t find an answer to those queries.",
                    confidence=0.0,
                    source='hybrid',
                    metadata={'sub_queries': len(results_list)}
                )

            return QueryResult(
                content=refined,
                confidence=min_conf,
                source='hybrid',
                metadata={'sub_queries': len(results_list)}
            )

        except Exception as e:
            self.logger.error(f"Error processing query: {e}")
            return QueryResult(
                content=f"An error occurred: {str(e)}",
                confidence=0.0,
                source='error',
                metadata={'error': str(e)}
            )


    def process_query_with_llamaindex_router(self, query: str) -> QueryResult:
        """
        Alternate method: use the LlamaIndex RouterRetriever if installed.
        This automatically decides between the doc index or SQL index.
        We still do guardrails/spell correction, but we skip the custom routes.
        """
        if not self.llamaindex_router_retriever:
            return QueryResult(
                content="LlamaIndex router is not configured or installed.",
                confidence=0.0,
                source='router_unavailable',
                metadata={}
            )
        try:
            self.logger.info(f"Processing query with LlamaIndex Router: {query}")

            # Guardrails
            violation = self.guardrail_checker.check(query)
            if violation:
                return QueryResult(
                    content="I’m sorry, but I can’t help with that request.",
                    confidence=1.0,
                    source='guardrail',
                    metadata={'violation': violation}
                )

            # Spell-correct
            corrected_query = self.spell_corrector.correct(query)

            # For this approach, we do not do sub-query decomposition;
            # we rely on LlamaIndex to route in one shot.
            query_engine = self.llamaindex_router_retriever.as_query_engine()
            response = query_engine.query(corrected_query)

            # LlamaIndex typically returns an object with a 'response' attribute
            refined = self.answer_refiner.refine(response.response)

            return QueryResult(
                content=refined,
                confidence=1.0,  # LlamaIndex doesn't produce numeric conf
                source='llamaindex_router',
                metadata={'raw_route_response': str(response)}
            )
        except Exception as e:
            self.logger.error(f"LlamaIndex router error: {e}")
            return QueryResult(
                content=f"Router error: {str(e)}",
                confidence=0.0,
                source='error',
                metadata={'error': str(e)}
            )

    ####################################################################
    # INTERNAL METHODS
    ####################################################################

    def _process_sub_query(self, query: str) -> QueryResult:
        """
        Process a single sub-query:
          1) Decide route (greeting, SQL, or RAG)
          2) Attempt retrieval
          3) Fallback if needed
        """
        try:
            # Check greeting first
            route_decision = self.smart_router.route_query(query)
            route = route_decision.route
            conf = route_decision.confidence

            self.logger.info(f"Sub-query: '{query}' -> route={route} (conf={conf:.2f})")

            if route == 'greeting':
                # Return a greeting
                return self._handle_greeting(query)

            if route == 'sql' and conf > self.config.SQL_CONFIDENCE_THRESHOLD:
                # Attempt SQL retrieval
                sql_result = self.sql_retriever.retrieve(query)
                if sql_result.confidence > 0 and sql_result.content.strip():
                    return sql_result

                # Fallback to RAG
                self.logger.info("SQL retrieval failed or empty, falling back to RAG.")
                rag_result = self.rag_retriever.retrieve(query)
                if rag_result.confidence > 0 and rag_result.content.strip():
                    return rag_result

                # If RAG also fails, return empty
                return QueryResult(
                    content="",
                    confidence=0.0,
                    source='none',
                    metadata={'error': 'SQL and RAG both failed'}
                )
            else:
                # Default to RAG if route='rag' or low confidence on SQL
                rag_result = self.rag_retriever.retrieve(query)
                if rag_result.confidence > 0 and rag_result.content.strip():
                    return rag_result

                return QueryResult(
                    content="",
                    confidence=0.0,
                    source='none',
                    metadata={'error': 'RAG failed'}
                )

        except Exception as e:
            self.logger.error(f"Sub-query error: {e}")
            return QueryResult(
                content="",
                confidence=0.0,
                source='error',
                metadata={'error': str(e)}
            )

    def _handle_greeting(self, query: str) -> QueryResult:
        """Respond to greeting queries like "hello," "thanks," etc."""
        q_lower = query.lower()
        g_patterns = self.smart_router.pattern_router.greeting_patterns

        if re.search(g_patterns['hello'], q_lower):
            return QueryResult(
                content="Hello there! How can I help you today?",
                confidence=1.0,
                source='greeting',
                metadata={'greeting_type': 'hello'}
            )
        elif re.search(g_patterns['goodbye'], q_lower):
            return QueryResult(
                content="Goodbye! Have a great day!",
                confidence=1.0,
                source='greeting',
                metadata={'greeting_type': 'goodbye'}
            )
        elif re.search(g_patterns['thanks'], q_lower):
            return QueryResult(
                content="You’re very welcome. Let me know if there’s anything else I can help with!",
                confidence=1.0,
                source='greeting',
                metadata={'greeting_type': 'thanks'}
            )
        elif re.search(g_patterns['wellbeing'], q_lower):
            return QueryResult(
                content="I’m doing well, thank you. How can I assist you today?",
                confidence=1.0,
                source='greeting',
                metadata={'greeting_type': 'wellbeing'}
            )
        else:
            return QueryResult(
                content="Hello! How can I help you?",
                confidence=1.0,
                source='greeting',
                metadata={'greeting_type': 'default'}
            )

    def _format_results(self, results: List[Tuple[str, QueryResult]]) -> str:
        """
        Format multiple sub-query results into a single combined answer for the user.
        """
        lines = []
        for sub_query, qres in results:
            if qres.confidence > 0 and qres.content.strip():
                lines.append(f"**Sub-query**: {sub_query}\n**Answer**: {qres.content}\n")
            else:
                lines.append(f"**Sub-query**: {sub_query}\n**Answer**: (No answer found.)\n")
        return "\n".join(lines)
