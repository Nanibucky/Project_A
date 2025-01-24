import os
import requests
import logging
import re
from functools import lru_cache
from concurrent.futures import ThreadPoolExecutor, as_completed
from langchain_community.utilities import SQLDatabase
from langchain_community.agent_toolkits import create_sql_agent
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from langchain_community.retrievers import BM25Retriever
from langchain.retrievers import EnsembleRetriever
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from llama_index.core import (
    VectorStoreIndex,
    SimpleDirectoryReader,
    StorageContext,
    SimpleKeywordTableIndex,
    Settings
)
from typing import List
from llama_index.core.schema import Node, NodeWithScore
from llama_index.core.tools import RetrieverTool
from llama_index.core.retrievers import RouterRetriever
from llama_index.core.selectors import PydanticSingleSelector
from llama_index.llms.openai import OpenAI
from langchain_openai import ChatOpenAI
import PyPDF2
import docx

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[logging.FileHandler("hybrid_agent.log"), logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

# Load API keys from environment variables
openai_api_key = "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"
deepseek_api_key = "sk-f56d50b9f71f4241a3de71f9ce03fb7b"

# Constants for text splitting
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200

# Initialize LangChain's ChatOpenAI for the SQL agent
langchain_llm = ChatOpenAI(
    model="gpt-3.5-turbo",
    temperature=0,
    api_key=openai_api_key
)

# Initialize llama_index's OpenAI for PydanticSingleSelector
llama_index_llm = OpenAI(
    model="gpt-3.5-turbo",
    temperature=0,
    api_key=openai_api_key
)

# Set the llama_index LLM in Settings
Settings.llm = llama_index_llm


class DeepSeekReasoner:
    """Wrapper class to interact with the DeepSeek-V3 model."""
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.api_url = "https://api.deepseek.com/v1/chat/completions"  # DeepSeek API endpoint

    @lru_cache(maxsize=100)  # Cache responses to avoid redundant API calls
    def reason(self, context: str, query: str) -> str:
        """Send a query and context to the DeepSeek-V3 model and return the reasoned response."""
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        data = {
            "model": "deepseek-chat",  # Use DeepSeek-V3
            "messages": [
                {"role": "system", "content": "You are a helpful assistant. Your task is to provide clear, concise, and accurate answers based on the given context."},
                {"role": "user", "content": f"Context: {context}\n\nQuestion: {query}"}
            ],
            "temperature": 0,
            "max_tokens": 500
        }
        try:
            response = requests.post(self.api_url, headers=headers, json=data, timeout=10)
            response.raise_for_status()
            return response.json()["choices"][0]["message"]["content"]
        except (requests.RequestException, KeyError, IndexError) as e:
            logger.error(f"Error communicating with DeepSeek: {e}")
            return "An error occurred while processing your request."


class DenseRetriever:
    """Creates a dense retriever using FAISS and OpenAI embeddings."""
    def __init__(self, openai_api_key: str, docs: List[Document]):
        self.openai_api_key = openai_api_key
        self.docs = docs

    @lru_cache(maxsize=1)  # Cache the dense store to avoid recomputing embeddings
    def build_dense_store(self):
        embeddings = OpenAIEmbeddings(openai_api_key=self.openai_api_key)
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP
        )
        split_docs = text_splitter.split_documents(self.docs)
        dense_store = FAISS.from_documents(split_docs, embeddings)
        return dense_store.as_retriever()


class SQLAgentRetriever:
    """Custom retriever to wrap the SQL agent and execute the SQL query directly."""
    def __init__(self, sql_agent, db):
        self.sql_agent = sql_agent
        self.db = db  # Add the database connection

    def execute_sql(self, query: str) -> List[dict]:
        """Execute the SQL query and return the results as a list of dictionaries."""
        try:
            results = self.db.run(query)  # Execute the query using the database connection
            return results  # Return the raw results
        except Exception as e:
            logger.error(f"Error executing SQL query: {e}")
            return []

    def retrieve(self, query: str) -> List[NodeWithScore]:
        """Retrieve and execute the SQL query, then return the answer based on results."""
        try:
            # Step 1: Run the SQL agent to generate the query
            sql_query = self.sql_agent.run(query)
            logger.info(f"Generated SQL Query: {sql_query}")  # Log the generated query

            # Step 2: Execute the query and fetch the results
            results = self.execute_sql(sql_query)

            if results:
                # Step 3: Format the results into a readable string
                formatted_results = self._format_results(results)

                # Step 4: Create a Node object for llama_index
                node = Node(
                    text=formatted_results,
                    metadata={"retriever": "sql_retriever"}  # Add metadata
                )

                # Step 5: Wrap the Node object in a NodeWithScore object
                node_with_score = NodeWithScore(node=node, score=1.0)  # Use a default score of 1.0
                return [node_with_score]  # Return the NodeWithScore object
            else:
                node = Node(
                    text="No results found for the query.",
                    metadata={"retriever": "sql_retriever"}  # Add metadata
                )
                node_with_score = NodeWithScore(node=node, score=1.0)  # Wrap the error Node in a NodeWithScore
                return [node_with_score]
        except Exception as e:
            logger.error(f"Error retrieving data from SQL agent: {e}")
            node = Node(
                text=f"Error executing the query: {e}",
                metadata={"retriever": "sql_retriever", "error": str(e)}  # Add error metadata
            )
            node_with_score = NodeWithScore(node=node, score=1.0)  # Wrap the error Node in a NodeWithScore
            return [node_with_score]

    def _format_results(self, results: List[dict]) -> str:
        """Format the SQL query results into a human-readable string."""
        formatted = ""
        for row in results:
            formatted += ", ".join([f"{key}: {value}" for key, value in row.items()]) + "\n"
        return formatted



class HybridAgent:
    def __init__(self, database_input: str, docs_directory: str, openai_api_key: str, deepseek_api_key: str, langchain_llm: ChatOpenAI):
        """Initialize the hybrid agent with database and document processing capabilities."""
        self.openai_api_key = openai_api_key
        self.deepseek_api_key = deepseek_api_key
        self.docs_directory = docs_directory
        self.langchain_llm = langchain_llm  # Pass the LangChain LLM object
        self.deepseek_reasoner = DeepSeekReasoner(api_key=deepseek_api_key)  # Initialize DeepSeek Reasoner

        # Initialize components
        self.db = self._init_database(database_input)
        self.sql_agent = self._create_sql_agent()
        self.docs = self._load_documents()
        self.hybrid_retriever = self._create_hybrid_retriever()
        self.router_retriever = self._create_router_retriever()

    def _init_database(self, db_input: str) -> SQLDatabase:
        """Initialize database connection."""
        if db_input.startswith("sqlite://") or db_input.startswith("postgresql://"):
            return SQLDatabase.from_uri(db_input)
        elif os.path.isfile(db_input):
            return SQLDatabase.from_uri(f"sqlite:///{db_input}")
        else:
            raise ValueError("Invalid database input. Please provide a valid URI or file path.")

    def _create_sql_agent(self):
        """Create SQL agent using LangChain's create_sql_agent."""
        return create_sql_agent(
            llm=self.langchain_llm,  # Use LangChain's LLM
            db=self.db,
            agent_type="openai-tools",
            verbose=True
        )

    @lru_cache(maxsize=1)  # Cache documents to avoid reloading
    def _load_documents(self) -> List[Document]:
        """Load documents from a directory, supporting .pdf and .docx formats."""
        docs = []
        for filename in os.listdir(self.docs_directory):
            file_path = os.path.join(self.docs_directory, filename)
            if filename.endswith('.pdf'):
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    content = ''
                    for page in reader.pages:
                        content += page.extract_text()
                    doc = Document(page_content=content, metadata={'source': filename, 'id': filename})
                    docs.append(doc)
            elif filename.endswith('.docx'):
                doc = docx.Document(file_path)
                content = ''
                for paragraph in doc.paragraphs:
                    content += paragraph.text
                doc = Document(page_content=content, metadata={'source': filename, 'id': filename})
                docs.append(doc)
            else:
                logger.warning(f"Unsupported file type: {filename}")
        return docs

    def _create_hybrid_retriever(self) -> EnsembleRetriever:
        """Create hybrid retriever combining dense and sparse retrievers."""
        # Initialize the sparse and dense retrievers
        sparse_retriever = BM25Retriever.from_documents(self.docs)
        dense_retriever = DenseRetriever(openai_api_key=self.openai_api_key, docs=self.docs).build_dense_store()

        # Combine retrievers with dynamic weights
        return EnsembleRetriever(
            retrievers=[sparse_retriever, dense_retriever],
            weights=[0.3, 0.7]  # Adjust weights as needed
        )

    def _create_router_retriever(self) -> RouterRetriever:
        """Create a RouterRetriever using LlamaIndex's PydanticSingleSelector."""
        # Wrap the SQL agent in a custom retriever
        sql_retriever = SQLAgentRetriever(self.sql_agent, self.db)  # Pass both sql_agent and db

        # Define the retriever tools
        retriever_tools = [
            RetrieverTool.from_defaults(
                retriever=sql_retriever,  # Use the custom SQL retriever
                name="sql_retriever",
                description="Useful for answering questions about structured data (e.g., sales, revenue, customers, orders).",
            ),
            RetrieverTool.from_defaults(
                retriever=self.hybrid_retriever,
                name="rag_retriever",
                description="Useful for answering questions about unstructured data (e.g., research papers, technical concepts, document content).",
            ),
        ]

        # Create the RouterRetriever
        return RouterRetriever(
            retriever_tools=retriever_tools,
            selector=PydanticSingleSelector.from_defaults(llm=Settings.llm),  # Use llama_index's LLM
        )

    def process_rag_query(self, query: str) -> str:
        """Process RAG queries using the hybrid retriever and OpenAI."""
        try:
            # Step 1: Use the hybrid retriever to retrieve relevant documents
            relevant_docs = self.hybrid_retriever.get_relevant_documents(query)

            if not relevant_docs:
                return "I don't know. The data is not available in the documents."

            # Step 2: Combine results into context
            context = "\n".join([doc.page_content for doc in relevant_docs])

            # Step 3: Use OpenAI to generate the response
            prompt = f"""Based on the following context, answer the question concisely and directly:

            Question: {query}

            Context: {context[:3000]}

            Rules:
            1. Stick to the context provided.
            2. If the context does not contain the answer, respond with "I don't know."
            3. Avoid adding irrelevant information.

            Answer:"""

            response = Settings.llm.complete(prompt)  # Use llama_index's LLM
            return response.text
        except Exception as e:
            logger.error(f"Error processing RAG query: {e}")
            return f"Error processing query: {str(e)}"

    def process_complex_query(self, query: str) -> str:
        """Process complex queries by splitting them and routing to appropriate models."""
        try:
            # Step 1: Use the RouterRetriever to retrieve the appropriate retriever
            retrieved_nodes = self.router_retriever.retrieve(query)
            logger.info(f"Retrieved Nodes: {retrieved_nodes}")  # Log the retrieved nodes for debugging

            if retrieved_nodes and retrieved_nodes[0].node.metadata.get("retriever") == "sql_retriever":
                # Use the SQL retriever to execute the query
                result = self.sql_agent.run(query)
                return result
            else:
                # Use the RAG retriever to process the query
                return self.process_rag_query(query)
        except Exception as e:
            logger.error(f"Error processing complex query: {e}")  # Log the error
            return f"Error processing complex query: {str(e)}"


# Example usage:
if __name__ == "__main__":
    database_input = "postgresql://postgres:NANIbucky%40662000@localhost:5432/chinook"  # Replace with your actual database
    docs_directory = "/Users/tharun/Desktop/Research_papers"  # Replace with your docs directory

    agent = HybridAgent(database_input, docs_directory, openai_api_key, deepseek_api_key, langchain_llm)

    # Test complex query
    complex_query = "who is the artist for The balls on the wall??"
    complex_result = agent.process_complex_query(complex_query)
    print(complex_result)

